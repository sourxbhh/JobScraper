"""Resume parsing + scoring.

- extract_text: pull plain text from .docx / .pdf / .md / .txt uploads.
- keyword_score: deterministic, instant baseline reusing the scraper's SKILL_KEYWORDS.
- run_llm_evaluation: background-task body that fills an LLM-reasoned evaluation.
"""
import io
import json
from datetime import datetime

from sqlalchemy.orm import Session

from models.job import Job
from models.resume import Resume, ResumeEvaluation
from models.settings import get_setting
from services.scraper_service import SKILL_KEYWORDS
from services import llm_service
import prompts

# Weight per skill tier, reused for both job-skill detection and coverage scoring.
_TIER_WEIGHT = {"high_value": 5, "medium_value": 3, "low_value": 1}


def extract_text(content: bytes, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(t for t in parts if t and t.strip())
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    # .md / .txt / unknown -> decode as text
    return content.decode("utf-8", errors="replace")


def _job_skills(description: str) -> list[tuple[str, int]]:
    """Skills (with weight) that the job description actually mentions."""
    desc = (description or "").lower()
    found = []
    for tier, kws in SKILL_KEYWORDS.items():
        w = _TIER_WEIGHT[tier]
        for kw in kws:
            if kw in desc:
                found.append((kw, w))
    return found


def keyword_score(resume_text: str, job: Job) -> dict:
    """Coverage of the job's required skills by the resume. Returns score + matched/missing."""
    resume = (resume_text or "").lower()
    job_skills = _job_skills(job.description or "")
    if not job_skills:
        return {"score": 0.0, "matched": [], "missing": []}

    total_w = sum(w for _, w in job_skills)
    matched, missing, got_w = [], [], 0
    for kw, w in job_skills:
        if kw in resume:
            matched.append(kw)
            got_w += w
        else:
            missing.append(kw)

    score = round(100.0 * got_w / total_w, 1) if total_w else 0.0
    return {"score": score, "matched": matched, "missing": missing}


def create_evaluation(db: Session, resume_id: int, job_id: int) -> ResumeEvaluation:
    """Create the eval row with the instant keyword score; LLM part fills in async."""
    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    job = db.query(Job).filter(Job.id == job_id).first()
    if not resume or not job:
        raise ValueError("Resume or job not found")

    kw = keyword_score(resume.raw_text, job)
    ev = ResumeEvaluation(
        resume_id=resume_id,
        job_id=job_id,
        status="pending",
        keyword_score=kw["score"],
        matched_json=json.dumps(kw["matched"]),
        missing_json=json.dumps(kw["missing"]),
    )
    db.add(ev)
    db.commit()
    db.refresh(ev)
    return ev


def run_llm_evaluation(db: Session, eval_id: int) -> None:
    """Background task: enrich an evaluation with LLM reasoning."""
    ev = db.query(ResumeEvaluation).filter(ResumeEvaluation.id == eval_id).first()
    if not ev:
        return
    ev.status = "running"
    db.commit()
    try:
        resume = db.query(Resume).filter(Resume.id == ev.resume_id).first()
        job = db.query(Job).filter(Job.id == ev.job_id).first()
        system, user = prompts.resume_eval(job, resume.raw_text)
        data = llm_service.chat_json(system, user)

        ev.llm_fit_score = float(data.get("fit_score", 0) or 0)
        feedback_parts = []
        if data.get("verdict"):
            feedback_parts.append(str(data["verdict"]))
        if data.get("strengths"):
            feedback_parts.append("Strengths: " + "; ".join(map(str, data["strengths"])))
        if data.get("gaps"):
            feedback_parts.append("Gaps: " + "; ".join(map(str, data["gaps"])))
        ev.llm_feedback = "\n".join(feedback_parts)
        ev.suggestions_json = json.dumps(data.get("suggestions", []))
        # Prefer the LLM's missing-keyword read if present; else keep keyword version.
        if data.get("missing_keywords"):
            ev.missing_json = json.dumps(data["missing_keywords"])
        ev.status = "done"
    except Exception as e:
        ev.status = "failed"
        ev.error = str(e)
    db.commit()


def primary_resume_text(db: Session) -> str | None:
    """CV context for other LLM features: explicit primary, else most recent resume."""
    pid = get_setting(db, "primary_resume_id")
    resume = None
    if pid:
        resume = db.query(Resume).filter(Resume.id == int(pid)).first()
    if not resume:
        resume = db.query(Resume).order_by(Resume.created_at.desc()).first()
    return resume.raw_text if resume else None
